from __future__ import annotations

import argparse
import html
import re
import shutil
from dataclasses import dataclass
from html import escape
from pathlib import Path

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from lxml import etree


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
DEFAULT_LAW = SKILL_DIR / "assets" / "law.png"

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
}

RISK_PROMPT = (
    "风险提示：本报告研究分析及结论完全基于公开数据进行测算和分析，相关模型构建与测算均基于"
    "国信金融工程团队客观研究。基金过往业绩及基金经理管理其他产品的历史业绩不代表未来表现，"
    "对基金产品和基金管理人的研究分析结论并不预示其未来表现，不能保证未来可持续性，亦不构成"
    "投资收益的保证或投资建议。本报告不涉及证券投资基金评价业务，不涉及对基金公司、基金经理、"
    "基金产品的推荐，亦不涉及对任何指数样本股的推荐。请详细阅读报告风险提示及声明部分。"
)

PROFILE_CARD = (
    '<mp-common-profile class="custom_select_card mp_profile_iframe mp_common_widget js_wx_tap_highlight" '
    'data-pluginname="mpprofile" data-nickname="量化藏经阁" data-alias="Z-quant" data-from="0" '
    'data-headimg="http://mmbiz.qpic.cn/mmbiz_png/ndYialDBEKPVyhzxrJzd8AFIfSQaiahaia1iaJJr6XMcheQthLicRewHTY7r5wZiaib1qURxRX5fCmt2AGtRXcxpqeo9g/300?wx_fmt=png&wxfrom=19" '
    'data-signature="分享量化投资和FOF投资领域的研究成果。" data-id="MzI5MzcwNTQ4NQ==" '
    'data-origin_num="1811" data-is_biz_ban="0" data-isban="0" data-biz_account_status="0" '
    'data-verify_status="0" data-index="0"></mp-common-profile>'
)

CN_NUMERALS = "一二三四五六七八九十"
BODY_LINE_HEIGHT = "1.6"
SUMMARY_BODY_STYLE_IDS = {"20", "23"}
SUMMARY_TEXT_STYLE = (
    f"margin:0;line-height:{BODY_LINE_HEIGHT};font-size:15px;color:#333333;"
    "text-align:justify;box-sizing:border-box;max-width:100%;overflow-wrap:break-word;"
)


@dataclass
class Metadata:
    title: str
    publication_date: str
    analysts: list[str]


def clean_text(value: str | None) -> str:
    if not value:
        return ""
    value = html.unescape(value)
    value = value.replace("\uf06c", "").replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def is_summary_body_style(style_values: list[str]) -> bool:
    return any(style in SUMMARY_BODY_STYLE_IDS for style in style_values)


def normalize_filename(value: str) -> str:
    value = re.sub(r'[\\/:*?"<>|\s]+', "-", value.strip())
    value = re.sub(r"-+", "-", value).strip("-")
    return value[:80] or "gs-report"


def iter_blocks(document: Document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def text_nodes(document: Document) -> list[str]:
    root = etree.fromstring(document.element.xml.encode("utf-8"))
    return [clean_text(node.text) for node in root.xpath(".//w:t", namespaces=NS) if clean_text(node.text)]


def paragraph_text(paragraph: Paragraph) -> str:
    return clean_text("".join(run.text for run in paragraph.runs) or paragraph.text)


def paragraph_is_bullet(paragraph: Paragraph) -> bool:
    numid = paragraph._p.xpath("./w:pPr/w:numPr/w:numId/@w:val")
    return bool(numid)


def clean_heading_title(value: str) -> str:
    return re.sub(r"^附录[一二三四五六七八九十]+[：:、.\s]*", "", value).strip() or value


def row_unique_texts(row) -> list[str]:
    texts: list[str] = []
    for cell in row.cells:
        text = clean_text(cell.text)
        if text and (not texts or texts[-1] != text):
            texts.append(text)
    return texts


def table_has_image(table: Table) -> bool:
    root = etree.fromstring(table._element.xml.encode("utf-8"))
    refs = root.xpath(".//a:blip/@r:embed | .//v:imagedata/@r:id", namespaces=NS)
    return bool(refs)


def table_title(table: Table) -> str:
    if not table.rows:
        return "图表"
    titles = [text for text in row_unique_texts(table.rows[0]) if not text.startswith("资料来源")]
    return titles[0] if titles else "图表"


def document_header_text(document: Document) -> str:
    values: list[str] = []
    for section in document.sections:
        for header in (section.header, section.first_page_header):
            values.extend(paragraph.text for paragraph in header.paragraphs)
    return "\n".join(values)


def extract_title(nodes: list[str]) -> str:
    if "核心观点" in nodes:
        end = nodes.index("核心观点")
    else:
        end = min(len(nodes), 8)
    candidates = [
        text
        for text in nodes[:end]
        if text
        and not text.startswith("国信证券")
        and "证券研究报告" not in text
        and "GUOSEN" not in text
        and text not in {"↵"}
    ]
    title = "".join(candidates).strip()
    return title or "未命名研究报告"


def extract_publication_date(document: Document, nodes: list[str], override: str | None) -> str:
    if override:
        return override
    haystack = document_header_text(document) + "\n" + "\n".join(nodes[:80])
    match = re.search(r"20\d{2}年\s*\d{1,2}月\s*\d{1,2}日", haystack)
    if not match:
        raise SystemExit("Missing publication date. Rerun with --publication-date YYYY年M月D日.")
    return re.sub(r"\s+", "", match.group(0))


def extract_analysts(nodes: list[str], overrides: list[str]) -> list[str]:
    if overrides:
        return overrides
    analysts: list[str] = []
    for index, text in enumerate(nodes):
        if text.startswith("证券分析师："):
            name = text.split("：", 1)[1].strip()
            cert = ""
            for lookahead in nodes[index + 1 : index + 12]:
                cert_match = re.search(r"S\d{13}", lookahead)
                if cert_match:
                    cert = cert_match.group(0)
                    break
            if name and cert:
                analysts.append(f"{name} {cert}")
    if not analysts:
        raise SystemExit('Missing analyst metadata. Rerun with --analyst "姓名 Sxxxxxxxxxxxxx".')
    deduped: list[str] = []
    for item in analysts:
        if item not in deduped:
            deduped.append(item)
    return deduped


def xml_run_is_bold(rnode) -> bool:
    bold_nodes = rnode.xpath("./w:rPr/w:b", namespaces=NS)
    if not bold_nodes:
        return False
    value = bold_nodes[0].get(f"{{{NS['w']}}}val")
    return value not in {"0", "false", "False", "off"}


def paragraph_node_html(pnode) -> str:
    parts: list[str] = []
    current_parts: list[str] = []
    current_bold: bool | None = None
    for rnode in pnode.xpath("./w:r", namespaces=NS):
        text = "".join(rnode.xpath(".//w:t/text()", namespaces=NS))
        if not text:
            continue
        is_bold = xml_run_is_bold(rnode)
        if current_bold is None:
            current_bold = is_bold
        if is_bold != current_bold:
            merged = escape("".join(current_parts))
            parts.append(f"<strong>{merged}</strong>" if current_bold else merged)
            current_parts = []
            current_bold = is_bold
        current_parts.append(text)
    if current_parts:
        merged = escape("".join(current_parts))
        parts.append(f"<strong>{merged}</strong>" if current_bold else merged)
    return "".join(parts)


def run_is_bold(run) -> bool:
    if run.bold is True:
        return True
    bold_nodes = run._r.xpath("./w:rPr/w:b")
    if not bold_nodes:
        return False
    value = bold_nodes[0].get(f"{{{NS['w']}}}val")
    return value not in {"0", "false", "False", "off"}


def paragraph_html(paragraph: Paragraph) -> str:
    parts: list[str] = []
    current_parts: list[str] = []
    current_bold: bool | None = None
    for run in paragraph.runs:
        if not run.text:
            continue
        is_bold = run_is_bold(run)
        if current_bold is None:
            current_bold = is_bold
        if is_bold != current_bold:
            merged = escape("".join(current_parts))
            parts.append(f"<strong>{merged}</strong>" if current_bold else merged)
            current_parts = []
            current_bold = is_bold
        current_parts.append(run.text)
    if current_parts:
        merged = escape("".join(current_parts))
        parts.append(f"<strong>{merged}</strong>" if current_bold else merged)
    return "".join(parts) or escape(paragraph_text(paragraph))


def is_heading_like_body_paragraph(text: str, content_html: str = "") -> bool:
    cleaned = clean_text(text)
    if not cleaned or len(cleaned) > 40:
        return False
    if re.match(r"^[一二三四五六七八九十]+、", cleaned):
        return True
    if cleaned.startswith(("（一）", "（二）", "（三）", "（四）", "（五）")):
        return True
    if content_html and re.fullmatch(r"(?:<strong>.*?</strong>)+", content_html):
        return True
    return False


def should_render_body_bullet(
    text: str,
    content_html: str,
    is_bullet: bool,
    previous_heading_like: bool = False,
) -> bool:
    if not is_bullet:
        return False
    if len(clean_text(text)) > 60:
        return False
    if is_heading_like_body_paragraph(text, content_html):
        return False
    if previous_heading_like:
        return False
    return True


def iter_table_paragraph_nodes(document: Document):
    for table in document.tables:
        root = etree.fromstring(table._element.xml.encode("utf-8"))
        for pnode in root.xpath(".//w:p", namespaces=NS):
            text = clean_text("".join(pnode.xpath(".//w:t/text()", namespaces=NS)))
            if not text:
                continue
            style = pnode.xpath("./w:pPr/w:pStyle/@w:val", namespaces=NS)
            numid = pnode.xpath("./w:pPr/w:numPr/w:numId/@w:val", namespaces=NS)
            yield text, bool(numid), style, paragraph_node_html(pnode)


def extract_summary(document: Document) -> list[tuple[str, list[tuple[str, bool, str]]]]:
    paragraphs = list(iter_table_paragraph_nodes(document))
    start = next((index for index, item in enumerate(paragraphs) if item[0] == "核心观点"), None)
    if start is None:
        raise SystemExit("Missing first-page 核心观点 area. Provide or add summary extraction support.")

    core_rows: list[tuple[str, bool, str]] = []
    in_core_body = False
    for text, is_bullet, style, inline_html in paragraphs[start + 1 :]:
        if text.startswith("风险提示："):
            break
        if is_summary_body_style(style):
            in_core_body = True
        if not in_core_body:
            continue
        core_rows.append((text, is_bullet, inline_html))

    if not core_rows or len(core_rows) <= 4:
        fallback_rows: list[tuple[str, bool, str]] = []
        after_related = False
        in_report_list = False
        for text, is_bullet, style, inline_html in paragraphs[start + 1 :]:
            if text.startswith("风险提示："):
                break
            if text == "相关研究报告":
                after_related = True
                in_report_list = True
                continue
            if not after_related:
                continue
            is_related_report_line = text.startswith("《") and "》" in text
            if in_report_list and is_related_report_line:
                continue
            if in_report_list and not is_related_report_line:
                in_report_list = False
            if not in_report_list:
                fallback_rows.append((text, is_bullet, inline_html))
        if fallback_rows:
            core_rows = fallback_rows

    groups: list[tuple[str, list[tuple[str, bool, str]]]] = []
    current_heading = ""
    current_items: list[tuple[str, bool, str]] = []
    for index, (text, is_bullet, inline_html) in enumerate(core_rows):
        next_is_body = index + 1 < len(core_rows)
        looks_heading = (
            not is_bullet
            and next_is_body
            and len(text) <= 45
            and not text.endswith(("。", "；", "，", "、"))
        )
        if looks_heading:
            if current_heading:
                groups.append((current_heading, current_items))
            current_heading = text
            current_items = []
        elif current_heading:
            current_items.append((text, is_bullet, inline_html))

    if current_heading:
        groups.append((current_heading, current_items))
    groups = [(heading, items) for heading, items in groups if items]
    if not groups and core_rows:
        return [("核心观点", core_rows)]
    if not groups:
        raise SystemExit("Failed to extract structured 核心观点 paragraphs.")
    return groups


def render_summary(groups: list[tuple[str, list[tuple[str, bool, str]]]]) -> str:
    parts = [
        '<section style="box-sizing:border-box;width:100%;max-width:100%;margin:12px 0 22px 0;padding:0 16px 18px 16px;border:1px solid #d6d6d6;border-radius:6px;background:#f7f7f7;overflow:hidden;">',
        '<section style="text-align:center;margin:0 0 12px 0;line-height:0;">'
        '<section style="display:inline-block;vertical-align:top;border-right:0.4em solid rgb(4,55,90);border-bottom:0.4em solid rgb(4,55,90);box-sizing:border-box;line-height:0;max-width:5%;border-top:0.4em solid transparent !important;border-left:0.4em solid transparent !important;"><br /></section>'
        '<section style="padding:2px 5px;display:inline-block;vertical-align:top;border-radius:0 0 5px 5px;background-color:rgb(5,68,119);font-size:17px;line-height:1.6;color:rgb(255,255,255);box-sizing:border-box;max-width:90%;"><p style="margin:0;box-sizing:border-box;"><strong style="font-size:16px;">&nbsp;&nbsp;报 告 摘 要 &nbsp;</strong></p></section>'
        '<section style="display:inline-block;vertical-align:top;border-bottom:0.4em solid rgb(4,55,90);border-left:0.4em solid rgb(4,55,90);box-sizing:border-box;line-height:0;max-width:5%;border-right:0.4em solid transparent !important;border-top:0.4em solid transparent !important;"><br /></section>'
        '</section>',
    ]
    for index, (heading, items) in enumerate(groups, 1):
        numeral = CN_NUMERALS[index - 1] if index <= len(CN_NUMERALS) else str(index)
        parts.append(
            f'<p style="margin:0;line-height:{BODY_LINE_HEIGHT};font-size:15px;box-sizing:border-box;max-width:100%;">'
            f'<span style="color:#c00000;font-weight:700;">{numeral}、{escape(heading)}</span></p>'
        )
        in_list = False
        for _, is_bullet, item_html in items:
            if is_bullet:
                if not in_list:
                    parts.append(
                        f'<ul style="margin:0;padding-left:1.2em;line-height:{BODY_LINE_HEIGHT};font-size:15px;'
                        'color:#333333;box-sizing:border-box;max-width:100%;overflow-wrap:break-word;">'
                    )
                    in_list = True
                parts.append(
                    f'<li data-gs-summary-bullet="true" style="{SUMMARY_TEXT_STYLE}">{item_html}</li>'
                )
            else:
                if in_list:
                    parts.append("</ul>")
                    in_list = False
                parts.append(f'<p style="{SUMMARY_TEXT_STYLE}">{item_html}</p>')
        if in_list:
            parts.append("</ul>")
    parts.append("</section>")
    return "\n\n".join(parts)


def render_h1(number: int, title: str) -> str:
    label = CN_NUMERALS[number - 1] if number <= len(CN_NUMERALS) else str(number)
    margin = "0 0 12px 0" if title == "总结" else "26px 0 12px 0"
    return f"""<section style="display:flex;flex-flow:row nowrap;margin:{margin};">
  <section style="padding:0 2px;display:inline-block;vertical-align:bottom;width:auto;flex:0 0 auto;align-self:flex-end;">
    <section style="margin-bottom:2px;"><span style="display:inline-block;background:#343d5d;color:#ffffff;font-weight:700;font-size:17px;line-height:32px;padding:0 22px;">{label}</span></section>
    <section style="background:#343d5d;height:3px;line-height:0;font-size:0;">&nbsp;</section>
  </section>
  <section style="display:inline-block;vertical-align:bottom;width:auto;flex:100 100 0%;align-self:flex-end;">
    <section style="padding:0 3px;color:#cc202a;letter-spacing:0;font-weight:700;font-size:16px;line-height:2;margin-bottom:2px;">{escape(title)}</section>
    <section style="background:#cc202a;height:4px;line-height:0;font-size:0;">&nbsp;</section>
  </section>
</section>"""


def render_h2(number: int, title: str) -> str:
    return f"""<section style="display:flex;flex-flow:row nowrap;margin:18px 0 8px 0;">
  <section style="padding:0 2px;display:inline-block;vertical-align:bottom;width:auto;flex:0 0 auto;align-self:flex-end;">
    <section style="margin-bottom:2px;text-align:center;"><span style="display:inline-block;color:#343d5d;font-weight:700;font-size:15px;line-height:1.6;padding:0 18px;">{number}</span></section>
    <section style="background:#343d5d;height:3px;line-height:0;font-size:0;">&nbsp;</section>
  </section>
  <section style="display:inline-block;vertical-align:bottom;width:auto;flex:100 100 0%;align-self:flex-end;">
    <section style="padding:0 3px;color:#cc202a;letter-spacing:0;font-weight:700;font-size:15px;line-height:1.8;margin-bottom:2px;">{escape(title)}</section>
    <section style="background:#cc202a;height:3px;line-height:0;font-size:0;">&nbsp;</section>
  </section>
</section>"""


def render_paragraph(content_html: str, is_bullet: bool = False) -> str:
    if is_bullet:
        return (
            f'<ul style="margin:0;padding-left:1.2em;line-height:{BODY_LINE_HEIGHT};font-size:15px;color:#333333;">'
            f'<li data-gs-body-bullet="true" style="margin:0;line-height:{BODY_LINE_HEIGHT};font-size:15px;'
            f'color:#333333;text-align:justify;">{content_html}</li></ul>'
        )
    return f'<p style="margin:0;line-height:{BODY_LINE_HEIGHT};font-size:15px;color:#333333;text-align:justify;">{content_html}</p>'


def render_marker(kind: str, number: int, title: str) -> str:
    return f'<p data-gs-marker="{kind}{number}" style="margin:0;line-height:{BODY_LINE_HEIGHT};font-size:14px;color:#333333;">{kind}{number}：{escape(title)}</p>'


def render_footer(metadata: Metadata, asset_rel: str) -> str:
    parts = [
        '<section style="height:10px;line-height:0;font-size:0;">&nbsp;</section>',
        f'<p style="margin:0;line-height:{BODY_LINE_HEIGHT};font-size:15px;font-weight:700;color:#333333;">注：本文选自国信证券于{escape(metadata.publication_date)}发布的研究报告《{escape(metadata.title)}》</p>',
    ]
    for analyst in metadata.analysts:
        parts.append(
            f'<p style="margin:0;line-height:{BODY_LINE_HEIGHT};font-size:15px;font-weight:700;color:#333333;">分析师：{escape(analyst)}</p>'
        )
    parts.extend(
        [
            f'<p style="margin:18px 0 0 0;"><img src="{escape(asset_rel)}" alt="法律声明" style="display:block;width:100%;height:auto;margin:0 auto;border:0;" /></p>',
        ]
    )
    return "\n\n".join(parts)


def build_markdown(document: Document, metadata: Metadata, asset_rel: str) -> str:
    nodes = text_nodes(document)
    groups = extract_summary(document)
    parts = [f"# {metadata.title}\n", render_summary(groups)]

    h1_count = 0
    h2_count = 0
    figure_count = 0
    table_count = 0
    started = False
    previous_body_heading_like = False
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = paragraph_text(block)
            if not text:
                continue
            style = block.style.name
            if style == "国信研报正文-1.正文一级标题":
                started = True
                h1_count += 1
                h2_count = 0
                previous_body_heading_like = False
                parts.append(render_h1(h1_count, clean_heading_title(text)))
            elif started and style == "国信研报正文-2.正文二级标题":
                h2_count += 1
                previous_body_heading_like = False
                parts.append(render_h2(h2_count, clean_heading_title(text)))
            elif started and style == "Normal" and text == "免责声明":
                break
            elif started and style in {"国信研报正文-4.正文", "Normal"}:
                content_html = paragraph_html(block)
                is_bullet = paragraph_is_bullet(block)
                heading_like = is_heading_like_body_paragraph(text, content_html)
                render_as_bullet = should_render_body_bullet(
                    text, content_html, is_bullet, previous_body_heading_like
                )
                parts.append(render_paragraph(content_html, render_as_bullet))
                previous_body_heading_like = heading_like
        elif started:
            previous_body_heading_like = False
            title = table_title(block)
            if table_has_image(block):
                image_refs = etree.fromstring(block._element.xml.encode("utf-8")).xpath(
                    ".//a:blip/@r:embed | .//v:imagedata/@r:id", namespaces=NS
                )
                for _ in image_refs:
                    figure_count += 1
                    parts.append(render_marker("图", figure_count, title))
            else:
                if title and len(block.rows) > 1:
                    table_count += 1
                    parts.append(render_marker("表", table_count, title))
    parts.append(render_footer(metadata, asset_rel))
    return "\n\n".join(part for part in parts if part) + "\n"


def write_outputs(markdown: str, output_dir: Path, slug: str, title: str) -> tuple[Path, Path]:
    md_path = output_dir / f"{slug}.md"
    html_path = output_dir / f"{slug}.html"
    md_path.write_text(markdown, encoding="utf-8")
    html_doc = f"""<!doctype html>
<html lang="zh-CN">
<head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{escape(title)}</title>
<style>
* {{ box-sizing: border-box; }}
body {{ max-width: 720px; margin: 24px auto; padding: 0 16px; font-family: -apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif; color:#333; }}
h1 {{ font-size: 22px; line-height: 1.35; margin: 0 0 12px; }}
img {{ max-width: 100%; }}
</style>
</head>
<body>
{markdown}
</body>
</html>
"""
    html_path.write_text(html_doc, encoding="utf-8")
    return md_path, html_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert a Guosen Word report into WeChat Markdown/HTML.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--output-dir", type=Path, default=Path.cwd())
    parser.add_argument("--template", type=Path, default=None, help="Accepted for workflow compatibility.")
    parser.add_argument("--slug", default=None)
    parser.add_argument("--publication-date", default=None)
    parser.add_argument("--analyst", action="append", default=[])
    parser.add_argument("--law-image", type=Path, default=DEFAULT_LAW)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    document = Document(args.docx)
    nodes = text_nodes(document)
    title = extract_title(nodes)
    publication_date = extract_publication_date(document, nodes, args.publication_date)
    analysts = extract_analysts(nodes, args.analyst)
    metadata = Metadata(title=title, publication_date=publication_date, analysts=analysts)

    args.output_dir.mkdir(parents=True, exist_ok=True)
    slug = args.slug or normalize_filename(title)
    asset_dir = args.output_dir / "assets" / slug
    asset_dir.mkdir(parents=True, exist_ok=True)
    if not args.law_image.exists():
        raise SystemExit(f"Missing law image: {args.law_image}")
    law_dest = asset_dir / "law.png"
    shutil.copyfile(args.law_image, law_dest)
    asset_rel = law_dest.relative_to(args.output_dir).as_posix()

    markdown = build_markdown(document, metadata, asset_rel)
    md_path, html_path = write_outputs(markdown, args.output_dir, slug, title)
    print(f"markdown={md_path}")
    print(f"html={html_path}")
    print(f"assets={asset_dir}")
    print(f"analysts={len(analysts)}")


if __name__ == "__main__":
    main()
